import pandas as pd
import docx
import json
import re
#from ..main import logger

#global text that will be added to the file
o_texte = ""
# this is the names of the columns in the csv file, they are removed.
# need to be translated in english
List_mot_a_enlever = ['nan', 'Annotation']

# the file mapper should be store at the root of the project
chemin_du_file_mapper = '../file_mapper.json'

# functions:

def add_highlighted_text(highlighted_text:str):
    global o_texte
    if (str(highlighted_text) not in List_mot_a_enlever):
        o_texte = o_texte + '\n\n' + str(highlighted_text)

# note management

def add_note(note:str):
    global o_texte
    if (str(note) not in List_mot_a_enlever):
        o_texte = o_texte + ' => ' + str(note)

def check_at(note:str)->bool:
    if '@@' in note:
        return True
    return False


def extract_keyword(text:str)-> str:
    # Remember: note are tagged with @@, the first word is the keyword, later used for the directory and then there is the comment
    # example: extract the keyword "sales" from the note "@@sales commentaire1 commentaire2"
    pattern = re.compile(r"""(?<=@@)([A-ù]+)(?= )""")
    match = re.search(pattern, text)
    if match:
        return (match[0].lower())
    #logger.error("error: keyword  not found")
    print("error: keyword  not found")


def extract_comment(text:str)-> str:
    # example: extract the comment "commentaire1 commentaire2" from the note  "##sales commentaire1 commentaire2"
    return (text.replace('@@' + extract_keyword(text), '=>'))

# creation of the output file

def add_to_file(filepath:str, surlignement:str, note:str):
    # add the note and the highlighted_text
    output_texte = surlignement + extract_comment(note)
    # append the text to the doc
    mydoc = docx.Document(filepath)
    mydoc.paragraphs[-1]
    mydoc.add_paragraph(output_texte)
    mydoc.save(filepath)

def process_csv_to_docx(chemin_du_csv:str, chemin_du_word:str):


    #1- always problems with separators of csv files ( virgules ou point virgules)- but the one of Amazon is standard
    #by default,  delimiter=',' and lineterminator (  backslash n)
    #initialization by loading a well configure csv file D.csv
    # the purpose of the try-catch

    df = pd.read_csv("C.csv")

    # 2- the interesting csv

    try:
        df = pd.read_csv(chemin_du_csv,sep=',')#, header=None,error_bad_lines='skip')
        # on_bad_lines= 'skip' for the new version of python ( not error_bad_lines)
    except pd.errors.ParserError: # pas sur de ce OSError
        #logger.error('cannot open')
        print('cannot open')
        #logger.error(
        print(
            """
from Amazon, the expected separator/delimitor: une virgule
expected qualifiers ( end of line):  retour à la ligne ( pas de ;) 
            """
        )


    with open(chemin_du_csv, encoding="utf8", errors='ignore') as fichier:
        contenu = fichier.read()

    # load file_mapper
    file_mapper = json.load(open(chemin_du_file_mapper))
    # rename columns "Type d'annotation","Emplacement","Suivi ?","Annotation"
    df.columns = ['Type', 'Emplacement', 'unknown', 'texte']

    for i in df.index:

        if df['Type'][i] == 'Surlignement (Jaune)':
            add_highlighted_text(df['texte'][i])

        elif df['Type'][i] == 'Note':
            if not check_at(df['texte'][i]):
                add_note(df['texte'][i])

            else:
                # management of the redirection of the note according to its title and the parameters store in the JSON of the file_mapper file.
                if extract_keyword(df['texte'][i]) in file_mapper.keys():
                    add_to_file(file_mapper[extract_keyword(df['texte'][i])], df['texte'][i - 1], df['texte'][i])

    mydoc = docx.Document()
    mydoc.add_paragraph(o_texte)
    mydoc.save(chemin_du_word)
    #logger.info("process_csv_to_docx finished")
    print("process_csv_to_docx finished")

def extract_Title( chemin_du_csv:str) -> str:
    return chemin_du_csv[:-4].replace("'"," ")
