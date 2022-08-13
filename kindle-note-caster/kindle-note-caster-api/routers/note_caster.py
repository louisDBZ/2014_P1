### import
import pandas as pd
import numpy as np
import docx
import json
import re
import csv
import logging


o_texte = ""
List_mot_a_enlever = ['nan', 'Annotation']  # ce sont en fait les titres des colonnes (cf amélioration possible)

# le file mapper est à aller chercher à l'extérieur du projet
chemin_du_file_mapper = '../file_mapper.json'
basefilepath = ""



### fonctions:

def ajouter_surlignement(surlignement:str):
    global o_texte
    if (str(surlignement) not in List_mot_a_enlever):
        o_texte = o_texte + '\n\n' + str(surlignement)


def ajouter_note(note:str):
    global o_texte
    if (str(note) not in List_mot_a_enlever):
        o_texte = o_texte + ' => ' + str(note)


def num_emplacement(emplacement:str):
    """not used"""
    return int(emplacement[12:])


def check_hashtag(note:str):
    if '@@' in note:
        return True
    return False


def extract_keyword(text:str):
    # example: sales de "##sales commentaire1 commentaire2"
    pattern = re.compile(r"""(?<=@@)([A-ù]+)(?= )""")
    match = re.search(pattern, text)
    if match:
        return (match[0].lower())
    print("error: keyword  not found")


def extract_note(text:str):
    # example: "commentaire1 commentaire2" de "##sales commentaire1 commentaire2"
    return (text.replace('@@' + extract_keyword(text), '=>'))


def ajouter_au_fichier(filepath:str, surlignement:str, note:str):
    # ajout de la basefilepath pour correspondre au changement de dossier
    filepath = basefilepath + filepath
    # ajout de la note et du surlignement
    # preprocessing du texte
    output_texte = surlignement + extract_note(note)
    # ajout à la fin du texte
    mydoc = docx.Document(filepath)
    mydoc.paragraphs[-1]
    mydoc.add_paragraph(output_texte)
    mydoc.save(filepath)


### script

def process_csv_to_docx(chemin_du_csv:str, chemin_du_word:str):
    #pour indenter plusieurs lignes: ctrl alt L et tab

    '''
     toujours des problèmes le choix des séparateurs ( virgules ou point virgules) avec les Csv
    ce que je ne comprends pas, il faut lancer un csv simple une première fois avant que cela fonctionne: le but du fichier D
    '''

    # 1- le csv "simple" mais inutile
    df = pd.read_csv("C.csv")

    # 2- le csv interessant
    # ici j'ai mis un try catch à cause des différents formats des csv

    try:
        df = pd.read_csv(chemin_du_csv,sep=',')#, header=None,error_bad_lines='skip')
        # on_bad_lines= 'skip' pour la nouvelle version de python
    except pd.errors.ParserError: # pas sur de ce OSError
        print('cannot open')
        # à faire passer en logging
        print(
            """
from Amazon, the expected separator/delimitor: une virgule
expected qualifiers ( end of line):  retour à la ligne ( pas de ;) 
            """
        )
        # à faire passer en logging

    # télécharger l'intelliji data viewer

    with open(chemin_du_csv, encoding="utf8", errors='ignore') as fichier:
        contenu = fichier.read()

    # load file_mapper
    file_mapper = json.load(open(chemin_du_file_mapper))
    # rename columns "Type d'annotation","Emplacement","Suivi ?","Annotation"
    df.columns = ['Type', 'Emplacement', 'unknown', 'texte']

    for i in df.index:

        if df['Type'][i] == 'Surlignement (Jaune)':
            ajouter_surlignement(df['texte'][i])

        elif df['Type'][i] == 'Note':
            # vérifie qu'il n'y a pas de ##

            if not check_hashtag(df['texte'][i]):
                ajouter_note(df['texte'][i])

            else:
                # il y a bien un marquage de rediretion de notes: il faut placer la note dans le bon fichier du dossier Lecture
                # ce mapping (le bon endroit des fichiers) est stocké dans un dictionnaire
                # celui ci est documenté dans le README
                # ajouter dans le fichier concerné la note + l'emplacement précédent

                if extract_keyword(df['texte'][i]) in file_mapper.keys():
                    ajouter_au_fichier(file_mapper[extract_keyword(df['texte'][i])], df['texte'][i - 1], df['texte'][i])

    # pas de else: les autres cas sont les résidus de texte en haut

    mydoc = docx.Document()
    mydoc.add_paragraph(o_texte)
    mydoc.save(chemin_du_word)
    # logging à la place de print
    # "operation completed" => mettre qch de plus explicite
    # la commande est alors à modifier
    # à faire
    #logging.info("process_csv_to_docx finished")
    print("process_csv_to_docx finished")



def extract_Title( chemin_du_csv:str) -> str:
    # se rappeler de cette façon de typer

    #with open(chemin_du_csv, newline='') as csvfile:
     #   csvreader = csv.reader(csvfile)
        # by default,  delimiter=',' et on a le bon lineterminator
        # print(csvreader.line_num, csvreader.__next__())
        #for row in reader:
         #   print(', '.join(row))
        #print(len(reader1))
        #title=reader[1]

    #print(title)

    return "title"
